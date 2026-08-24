#!/usr/bin/env python3
"""cms.docx_export — one Word workbook per page in the style of the original
"GTS Website Content Template".

Each .docx follows the same layout as the reference workbook:
  * title + short how-to
  * sections named after what visitors see
  * card groups within sections
  * one table per section/card with rows of:
        Field | Current copy | Your new text | Max

"Current copy" is filled with the published CMS value (falling back to the
seed), "Your new text" is left blank for edits, and "Max" shows the design
ceiling as characters / words.
"""
import os
import json
import re
import sys
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import gts_cms

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.join(gts_cms.SITE, 'CMS Word Workbooks')
NAVY = RGBColor(0x1B, 0x2A, 0x41)
GOLD = RGBColor(0x8A, 0x7F, 0x3F)
GREY = RGBColor(0x66, 0x66, 0x66)

FIELD_LABELS = {'h1': 'Heading', 'h2': 'Heading', 'h3': 'Heading',
                'p': 'Paragraph', 'a': 'Link', 'blockquote': 'Quote',
                'figcaption': 'Attribution', 'li': 'Bullet', 'span': 'Inline text'}


def humanize(slug):
    return slug.replace('_', ' ').replace('-', ' ').strip()


def max_cell(spec):
    return '%dc / %dw' % (spec['max_chars'], spec['max_words'])


def current_value(key):
    return published.get(key, gts_cms.SEED.get(key, ''))


def para(doc, text, size=10, bold=False, italic=False, color=None,
         space_after=4, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, rows):
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(0.9), Inches(3.1), Inches(2.2), Inches(0.8)]
    hdr = t.rows[0].cells
    for i, name in enumerate(['Field', 'Current copy', 'Your new text', 'Max']):
        cell = hdr[i]
        cell.width = widths[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1B2A41')
        cell._tc.get_or_add_tcPr().append(shd)
    for label, cur, key in rows:
        cells = t.add_row().cells
        for i, val in enumerate([FIELD_LABELS.get(label, label), cur, '', max_cell(gts_cms.SCHEMA[key])]):
            cells[i].width = widths[i]
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
        cells[2].text = ''
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F5F0DC')
        cells[2]._tc.get_or_add_tcPr().append(shd)
    return t


def section_title(section_slug, keys):
    """Use the section's own heading copy as its display title."""
    if section_slug == 'marquee':
        return 'Marquee banner'
    best = ''
    for k in keys:
        parts = k.split('.')
        if parts[-1] in ('heading', 'title') and gts_cms.SCHEMA[k]['label'] in ('h1', 'h2'):
            v = current_value(k)
            if v:
                return v.strip()
        if parts[-1] == 'heading' and gts_cms.SCHEMA[k]['label'] == 'h3':
            v = current_value(k)
            if v and not best:
                best = v.strip()
    return best or humanize(section_slug)


GENERIC = {'kicker', 'heading', 'title', 'paragraph', 'link', 'quote',
           'attribution', 'bullet', 'accent', 'item', 'page'}


def build_sections(page_slug):
    """Group page keys into (section, [(card_title, rows)]) preserving order.

    Keys like page.paragraph.1 (no real section slug) belong to whichever
    section precedes them, so they are attached to that section instead of
    spawning bogus "paragraph"/"accent" groups. FAQ question/answer sets are
    split into individual cards titled by the question itself."""
    keys = [k for k in gts_cms.page_keys(page_slug)]
    groups = []  # list of {'slug','flat':[], 'cards':{card_title:[keys]}}
    for k in keys:
        parts = k.split('.')
        sec = parts[1] if len(parts) > 1 else 'page'
        card = None
        if len(parts) > 2 and parts[2] not in GENERIC and not parts[2].isdigit():
            card = parts[2]
        if sec in GENERIC:
            if groups:
                cur = groups[-1]
            else:
                groups.append({'slug': 'page', 'flat': [], 'cards': {}})
                cur = groups[-1]
        else:
            if not groups or groups[-1]['slug'] != sec:
                groups.append({'slug': sec, 'flat': [], 'cards': {}})
            cur = groups[-1]
        if card:
            cur['cards'].setdefault(card, []).append(k)
        else:
            cur['flat'].append(k)
    sections = []
    for g in groups:
        sec = g['slug']
        flat = g['flat']
        cards = g['cards']
        sec_keys = flat + [k for c in cards for k in cards[c]]
        title = section_title(sec, sec_keys)
        card_blocks = []
        for card in cards:
            rows = [(gts_cms.SCHEMA[k]['label'], current_value(k), k) for k in cards[card]]
            if card == 'question' and rows:
                # split each question + its answer into its own card, titled
                # by the question's current copy
                i = 0
                while i < len(rows):
                    q = rows[i]
                    if q[0] == 'summary':
                        j = i + 1
                        while j < len(rows) and rows[j][0] not in ('summary',):
                            j += 1
                        qtext = q[1].strip() or 'Question'
                        card_blocks.append((qtext, rows[i:j]))
                        i = j
                    else:
                        card_blocks.append(('', [q]))
                        i += 1
            else:
                card_blocks.append((humanize(card), rows))
        sections.append((title, flat, card_blocks, sec_keys))
    return sections


def make_docx(page, page_slug):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    para(doc, 'Grace & Truth Theological Seminary', size=16, bold=True, color=NAVY)
    para(doc, 'Website content workbook — %s PAGE' % page.upper(), size=12, bold=True, color=GOLD)
    para(doc, 'How to use: each row is one editable snippet. "Current copy" shows what is '
              'live on the site; write your replacement in the shaded "Your new text" cell next '
              'to it. The "Max" column is the design ceiling — target 60–80%. Buttons and links '
              'must stay on one line (~3 words). Keep the field types unchanged.',
         size=9, italic=True, color=GREY, space_after=10)

    sections = build_sections(page_slug)
    for idx, (title, flat, cards, sec_keys) in enumerate(sections, start=1):
        para(doc, 'Section %d — %s' % (idx, title), size=12, bold=True, color=NAVY, space_after=6)
        if flat:
            rows = [(gts_cms.SCHEMA[k]['label'], current_value(k), k) for k in flat]
            add_table(doc, rows)
            para(doc, '', size=4, space_after=2)
        for ci, (card_title, rows) in enumerate(cards, start=1):
            if len(cards) > 1 or card_title:
                label = 'Card %d — %s' % (ci, card_title)
            else:
                label = 'Card %d' % ci
            para(doc, label, size=10, bold=True, color=GOLD, space_after=4)
            add_table(doc, rows)
            para(doc, '', size=4, space_after=2)

    return doc


def main():
    os.makedirs(OUT, exist_ok=True)
    global published
    published = gts_cms.published_values(gts_cms.db_connect())
    for page in gts_cms.PAGES:
        slug = gts_cms.PAGE_SLUG[page]
        label = gts_cms.PAGE_LABEL[page]
        doc = make_docx(label, slug)
        path = os.path.join(OUT, 'GTS %s Content.docx' % label)
        doc.save(path)
        n = len(gts_cms.page_keys(slug))
        print('wrote %-10s (%2d fields) -> %s' % (label, n, path))


if __name__ == '__main__':
    main()
